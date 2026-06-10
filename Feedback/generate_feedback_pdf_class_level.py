import pandas as pd
import numpy as np
import mysql.connector
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import pypandoc, platform
pypandoc.download_pandoc()
import pandoc
import os
from pathlib import Path
from pypdf import PdfWriter, PdfReader
import re
from docx.shared import Inches

def safe_filename(name: str, max_len: int = 80) -> str:
    """
    Make string safe for Windows filenames.
    - Removes all special characters
    - Normalizes spaces
    - Trims length
    """
    if name is None:
        return "NA"

    name = str(name)

    # Replace all non safe chars with underscore
    name = re.sub(r'[^\w\-]+', '_', name)

    # Remove leading/trailing underscores
    name = name.strip('_')

    # Collapse multiple underscores
    name = re.sub(r'_+', '_', name)

    return name[:max_len]


def safe_convert(input_path, output_path):
    try:
        if platform.system() == "Windows":
            convert(input_path, output_path)
            print(f"Converted with docx2pdf: {input_path} → {output_path}")
        else:
            raise Exception("docx2pdf not supported on this OS")
    except Exception as e:
        print("docx2pdf failed, falling back to pypandoc:", e)
        try:
            pypandoc.convert_file(
                input_path,
                'pdf',
                outputfile=str(output_path),
                extra_args=['--standalone', '--pdf-engine=wkhtmltopdf']
            )
            print(f"Converted with wkhtmltopdf: {input_path} → {output_path}")
        except Exception as e2:
            print("Both conversions failed:", e2)

DB_NAME = "collpoll_micms"
# output folder for this run: "<schema suffix>_<session id>" (set in __main__)
RUN_FOLDER = None


def run_root():
    """Base output folder for this run, e.g. ~/Downloads/micms_51."""
    return Path.home() / "Downloads" / RUN_FOLDER


def folder_check(folder_name):
    folder_path = run_root() / folder_name
    if not os.path.exists(folder_path):
    # If it doesn't exist, create the folder
        os.makedirs(folder_path)
        print(f"Folder '{folder_path}' created.")
    else:
        print(f"Folder '{folder_path}' already exists.")

    return folder_path


def fetch_data(query):
    mydb = mysql.connector.connect(
        host="collpolldb9-read.c5sc77nejhmr.ap-south-1.rds.amazonaws.com",
        user="suraj_shetty",
        passwd="3qIGaWCdlh",
        database=DB_NAME
    )

    mycursor = mydb.cursor(dictionary=True)
    mycursor.execute(query)
    raw_data = mycursor.fetchall()
    return pd.DataFrame(raw_data)

from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove any existing width (python-docx adds an equal-split one by default,
    # and the first w:tcW wins) so our value actually takes effect
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcPr.append(tcW)


def set_grid_widths(table, widths):
    """Set the table's column grid (w:gridCol) so fixed-layout tables honour
    per-column widths instead of splitting evenly."""
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        return
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(w * 1440)))


def writeDfToDoc(df, docx):

    rows, cols = df.shape
    table = docx.add_table(rows + 1, cols)
    table.style = 'Table Grid'

    # 🔑 ABSOLUTELY REQUIRED
    table.autofit = False
    table.allow_autofit = False

    # ---- HEADER ----
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        run = cell.paragraphs[0].add_run(col)
        run.bold = True

    # ---- DATA ----
    for i in range(rows):
        table.cell(i + 1, 0).text = str(df.iloc[i, 0])
        table.cell(i + 1, 1).text = str(df.iloc[i, 1])

    # 🔥 FORCE COLUMN WIDTHS (CELL LEVEL)
    for row in table.rows:
        set_cell_width(row.cells[0], 1)   # S. No (narrow)
        set_cell_width(row.cells[1], 5.8)   # Student Response (wide)

        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def build_report_name(prefix, faculty, course_code, class_name):
    return f"{prefix}-" \
           f"{safe_filename(faculty)}_" \
           f"{safe_filename(course_code[:13])}_" \
           f"{safe_filename(class_name)}"


def createSubjectiveDocument(dataSub, faculty, course, term, course_code, class_name, saveLoc):
    SubjectiveDoc = Document()

    sec_pr = SubjectiveDoc.sections[0]._sectPr # get the section properties el
    # create new borders el
    pg_borders = OxmlElement('w:pgBorders')
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section
    # my_image = SubjectiveDoc.add_picture('MICA_Logo.png', width=Inches(1), height=Inches(1)) 
    # last_paragraph = SubjectiveDoc.paragraphs[-1] 
    # last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # last_paragraph.paragraph_format.space_before = Pt(1)
    
    question_text = dataSub['question_text'].drop_duplicates().to_list()
    p = SubjectiveDoc.add_paragraph()
    p.add_run('Faculty Name: ').bold = True
    p.add_run(faculty)
    p.add_run('\nCourse Name ').bold = True
    p.add_run(course)
    p.add_run('\nTerm: ').bold = True
    p.add_run(term)
    p.add_run('\nClass Name: ').bold = True
    p.add_run(class_name)
    
    for every in question_text:
        p = SubjectiveDoc.add_paragraph()
        p.add_run('\nQuestion: ').bold = True
        p.add_run(every)

        subset = dataSub[dataSub['question_text'] == every]
        sub = subset[['response_text']].reset_index(drop=True)
        sub.columns = ['Student Response']
        sub['S. No'] = ['Student ' + str(x + 1) for x in sub.index.values.tolist()]
        sub = sub[['S. No', 'Student Response']]

        writeDfToDoc(df=sub, docx=SubjectiveDoc)

    safe_code = safe_filename(course_code[0:13])
    fname = build_report_name(
        "sub",
        faculty,
        course_code,
        class_name
    )


    SubjectiveDoc.save(saveLoc / f"{fname}.docx")
    safe_convert(saveLoc / f"{fname}.docx", saveLoc / f"{fname}.pdf")
    os.remove(saveLoc / f"{fname}.docx")


OBJECTIVE_QTYPES = ("SINGLE_SELECT", "MULTI_SELECT")


def add_page_border(doc):
    """Apply a single-line page border to the first section (matches the old templates)."""
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el)
    sec_pr.append(pg_borders)


def parse_template_structure(struct_df):
    """Turn the template-definition rows (one row per option) into an ordered
    structure for ONE template.

    The feeding query (per template) is::

        select t4.id template_id, t4.template_name, t3.id section_id, t3.section_name,
               t2.id question_id, t2.question_text, t2.question_type,
               t1.option_text, t1.option_score
        from feedback_template t4
        left join feedback_template_section t3 on t3.feedback_template_id = t4.id
        left join feedback_template_section_question t2 on t2.template_section_id = t3.id
        left join feedback_template_question_option t1 on t1.feedback_question_id = t2.id
        order by t4.id, t3.id, t2.id, t1.option_score

    Returns a dict:
        sections          ordered [{section_name, questions:[{question_id, question_text,
                          options:[(text, score)...] sorted by score, position}]}]
        obj_question_ids  objective question ids in display order (drives itemRating{n})
        scales            distinct answer scales (ordered option_text tuples) for the
                          multi-scale header rows
        n_cols            number of option columns (max options across questions)
    """
    df = struct_df[struct_df['question_type'].isin(OBJECTIVE_QTYPES)].copy()
    df = df[df['option_text'].notna()]
    df['option_text'] = df['option_text'].astype(str).str.strip()
    df['section_name'] = df['section_name'].fillna('').astype(str).str.strip()

    sections = []
    obj_question_ids = []
    scales = []
    n_cols = 0
    pos = 0

    for _, sec_df in df.groupby('section_id', sort=False):
        questions = []
        for q_id, q_df in sec_df.groupby('question_id', sort=False):
            q_df = q_df.sort_values('option_score')
            opts = list(zip(q_df['option_text'].tolist(), q_df['option_score'].tolist()))
            pos += 1
            obj_question_ids.append(q_id)
            scale = tuple(t for t, _ in opts)
            if scale and scale not in scales:
                scales.append(scale)
            n_cols = max(n_cols, len(opts))
            questions.append({
                'question_id': q_id,
                'question_text': q_df['question_text'].iloc[0],
                'options': opts,
                'position': pos,
            })
        sections.append({
            'section_name': sec_df['section_name'].iloc[0],
            'questions': questions,
        })

    return {
        'sections': sections,
        'obj_question_ids': obj_question_ids,
        'scales': scales,
        'n_cols': n_cols,
    }


def _val(data, key):
    """Safe string lookup from a finalDataset row (blank for missing / NaN)."""
    if key in data:
        v = data[key]
        if pd.notna(v):
            return str(v)
    return ''


def _bold_cells(cells):
    for c in cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True


def _grid_row(table, values, bold=False):
    cells = table.add_row().cells
    for i, v in enumerate(values):
        cells[i].text = '' if v is None else str(v)
    if bold:
        _bold_cells(cells)
    return cells


def _avg_row(table, label, value, bold=True):
    """A row whose label spans every column except the last (the AVG value)."""
    row = table.add_row()
    row.height = Pt(26)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cells = row.cells
    merged = cells[0]
    for c in cells[1:-1]:
        merged = merged.merge(c)
    merged.text = label
    cells[-1].text = '' if value is None else str(value)
    if bold:
        _bold_cells([merged, cells[-1]])
    return cells


def _merge_down(cells, text, bold=True):
    """Vertically merge a column of cells into one and write `text` once."""
    top = cells[0]
    for c in cells[1:]:
        top = top.merge(c)
    top.text = text
    top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bold:
        _bold_cells([top])
    return top


# usable page width = letter (8.5") minus the 0.6" left/right margins set above
PAGE_USABLE_IN = 7.3


def _style_report_table(table, label_width=3.1, avg_width=0.55, usable=PAGE_USABLE_IN):
    """Small font everywhere, a wide first (question/label) column, the option
    columns sharing the remaining width, and all data cells centre-aligned so
    the table fills the page instead of leaving big side margins."""
    ncol = len(table.columns)
    nopt = ncol - 2  # columns between the label and the trailing AVG/value
    if nopt > 0:
        opt_w = max((usable - label_width - avg_width) / nopt, 0.4)
        widths = [label_width] + [opt_w] * nopt + [avg_width]
    else:
        widths = [label_width, usable - label_width]

    table.autofit = False
    table.allow_autofit = False
    set_grid_widths(table, widths)
    for row in table.rows:
        # give every normal row a bit more height (keep the taller average rows)
        if row.height is None:
            row.height = Pt(20)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                if ci > 0:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ci < len(widths):
                set_cell_width(cell, widths[ci])


def build_objective_document(tstruct, data, template_name, logo_path, save_loc):
    """Build the objective feedback report directly from the DB template
    structure (no .docx template / find-replace)."""
    doc = Document()
    add_page_border(doc)

    # narrower margins so the tables can use most of the page width
    for s in doc.sections:
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)
        s.top_margin = Inches(0.4)

    if logo_path and os.path.exists(logo_path):
        # reuse the document's first (empty) paragraph so there is no blank line
        # above the logo, and keep the image at its natural size / aspect ratio
        p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.add_run().add_picture(logo_path)

    # ---- header info block: two balanced columns, bold labels, no borders ----
    def field(cell, label, value):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        vr = p.add_run(value)
        vr.font.size = Pt(9)

    course = f"{_val(data, 'course_name')} ({_val(data, 'course_code')})"
    left = [
        ("Course Name: ", course),
        ("Faculty: ", _val(data, 'faculty_name')),
        ("Total # of students: ", _val(data, 'studentPresent')),
        ("No. of respondents: ", _val(data, 'totalResponse')),
    ]
    right = [
        ("Term: ", _val(data, 'term')),
        ("Course Type: ", _val(data, 'course_type')),
        ("Class: ", _val(data, 'class_name')),
        (None, None),
    ]

    hdr = doc.add_table(rows=len(left), cols=2)
    hdr.autofit = False
    hdr.allow_autofit = False
    for i in range(len(left)):
        field(hdr.cell(i, 0), *left[i])
        if right[i][0] is not None:
            field(hdr.cell(i, 1), *right[i])
    for row in hdr.rows:
        set_cell_width(row.cells[0], 3.7)
        set_cell_width(row.cells[1], 3.1)

    # minimal separator before the first table. Two tables need a paragraph
    # between them (else Word merges them), so force it to an exact tiny height.
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(0)
    gap.paragraph_format.line_spacing = Pt(2)
    gap.add_run().font.size = Pt(2)

    sections = tstruct['sections']
    multi = len(sections) > 1

    # ---- one table per section (so sections are visually separated), each with
    #      its own header: section name (merged) + its option labels + numbers ----
    for k, sec in enumerate(sections, start=1):
        corner = sec['section_name'] or template_name

        # distinct answer scales used within THIS section + its option count
        sec_scales = []
        sec_ncols = 0
        for q in sec['questions']:
            scale = tuple(t for t, _ in q['options'])
            if scale and scale not in sec_scales:
                sec_scales.append(scale)
            sec_ncols = max(sec_ncols, len(q['options']))

        table = doc.add_table(rows=0, cols=sec_ncols + 2)
        table.style = 'Table Grid'

        # header: one row per distinct scale, then the numeric row
        header_rows = []
        for scale in sec_scales:
            header_rows.append(_grid_row(table, [corner] + [scale[i] if i < len(scale) else '' for i in range(sec_ncols)] + ['AVG'], bold=True))
        header_rows.append(_grid_row(table, [corner] + [str(i + 1) for i in range(sec_ncols)] + ['AVG'], bold=True))

        # merge the repeated section-name corner (and the AVG header) vertically
        if len(header_rows) > 1:
            _merge_down([hr[0] for hr in header_rows], corner)
            _merge_down([hr[-1] for hr in header_rows], 'AVG')

        # question rows
        for q in sec['questions']:
            pos = q['position']
            opt_vals = [
                _val(data, f"Q{pos}_{q['options'][i][0]}") if i < len(q['options']) else ''
                for i in range(sec_ncols)
            ]
            _grid_row(table, [q['question_text']] + opt_vals + [_val(data, f"itemRating{pos}")])

        # closing average row for this section
        if multi:
            _avg_row(table, 'Section Average', _val(data, f'sec{k}'))
        else:
            _avg_row(table, 'Overall Average', _val(data, 'final'))

        _style_report_table(table)
        doc.add_paragraph()  # gap after each section

    # overall total across sections (only meaningful when there is more than one)
    if multi:
        total = doc.add_table(rows=0, cols=2)
        total.style = 'Table Grid'
        _avg_row(total, 'Total Average', _val(data, 'section_final') or _val(data, 'final'))
        _style_report_table(total, label_width=5.5)

    fname = build_report_name("obj", _val(data, 'faculty_name'),
                              _val(data, 'course_code'), _val(data, 'class_name'))
    doc.save(save_loc / f"{fname}.docx")
    safe_convert(save_loc / f"{fname}.docx", save_loc / f"{fname}.pdf")
    os.remove(save_loc / f"{fname}.docx")

    


if __name__ == "__main__":

    x = int(input("Enter Session Id: "))
    # output folder = "<schema suffix>_<session id>", e.g. collpoll_micms -> micms_51
    RUN_FOLDER = f"{DB_NAME.split('_')[-1]}_{x}"
    # single logo for every report generated in this run (top of page 1); blank = none
    logo_path = input("Enter logo image path (leave blank for none): ").strip().strip('"').strip("'") or None
    # achor_teaching = Document(r"C:\Users\Suraj Shetty\Downloads\PGP Level 2 - Anchor + Teaching - MICA FORMAT.docx")
    # teaching = Document(r"C:\Users\Suraj Shetty\Downloads\PGP-Level 2 (Teaching) - MICA FORMAT.docx")

    response = fetch_data('''
        Select -- distinct t1.student_ukid
    t1.*, t2.responded, t3.toRespond
    from (
        Select 
        fs.session_name, cl.id class_id, cl.batch as class_name, concat(ua.f_name, ' ', ua.l_name) faculty_name, fp.ukid faculty_ukid, 
        pr.programme_name, c.course_id, c.course_code, c.course_name, cl.type as course_type, term.name as term, term.id term_id,
        ft.template_name, fsr.question_id, fts.section_name, ftsq.question_text, fss.student_ukid, ftqo.option_text, ftqo.option_score,
        fcft.session_id, fcft.template_id
        from feedback_student_response fsr
        left join feedback_student_template_status fsts on fsts.id = fsr.session_student_template_status_id
        left join feedback_template_question_option ftqo on ftqo.id = fsr.option_id
        left join feedback_template_section_question ftsq on ftsq.id = fsr.question_id
        left join feedback_template_section fts on fts.id = ftsq.template_section_id
        left join feedback_session_student fss on fsts.student_session_id = fss.id
        left join feedback_session fs on fs.id = fss.session_id
        left join feedback_course_faculty_template fcft on fcft.id = fsts.feedback_course_faculty_template_id
        left join class cl on cl.id = fcft.class_id
        left join faculty_profile fp on fp.ukid = fcft.faculty_ukid
        left join user_attributes ua on fp.ukid = ua.ukid
        left join feedback_template ft on ft.id = fcft.template_id
        left join course c on c.course_id = cl.course_id
        left join programme pr on pr.programme_id = cl.programme_id
        left join term on term.id = cl.term_id
        where fsr.option_id is not null
        and fss.session_id = %d and fsts.submitted = 1
        group by question_id, fp.ukid, fs.session_name, ft.template_name, cl.batch, c.course_code, c.course_name, cl.type, student_ukid, ftqo.id
    ) t1
    left join (
    -- no of questions responded by each student
        Select fcft.session_id, faculty_ukid, fcft.class_id, template_id, student_ukid, count(*) responded
        from feedback_student_response fsr
        left join feedback_student_template_status fsts on fsts.id = fsr.session_student_template_status_id
        left join feedback_course_faculty_template fcft on fcft.id = fsts.feedback_course_faculty_template_id
        left join feedback_session_student fss on fsts.student_session_id = fss.id
        where fsr.option_id is not null
        group by fcft.session_id, faculty_ukid, template_id, student_ukid, fcft.class_id
    ) t2 on t1.session_id = t2.session_id 
        and t1.faculty_ukid = t2.faculty_ukid 
        and t1.template_id = t2.template_id 
        and t1.student_ukid = t2.student_ukid 
        and t1.class_id = t2.class_id
    left join (
    -- no of question in template
        Select feedback_template_id, count(*) toRespond
        from feedback_template_section_question ftsq
        left join feedback_template_section fts on fts.id = ftsq.template_section_id
        where question_type in  ("SINGLE_SELECT","MULTI_SELECT") 
        group by feedback_template_id
    ) t3 on t1.template_id = t3.feedback_template_id
    where responded > 0;
    ''' %x)

    response_subjective = fetch_data('''
    Select fs.session_name, cl.id class_id, cl.batch as class_name, concat(ua.f_name, ' ', ua.l_name) faculty_name, fp.ukid faculty_ukid, 
    pr.programme_name, c.course_code, c.course_name, cl.type as course_type, term.id term_id, term.name as term, fsr.response_text,
    ft.template_name, fsr.question_id, fts.section_name, ftsq.question_text, fss.student_ukid, ftqo.option_text, ftqo.option_score
    from feedback_student_response fsr
    left join feedback_student_template_status fsts on fsts.id = fsr.session_student_template_status_id
    left join feedback_template_question_option ftqo on ftqo.id = fsr.option_id
    left join feedback_template_section_question ftsq on ftsq.id = fsr.question_id
    left join feedback_template_section fts on fts.id = ftsq.template_section_id
    left join feedback_session_student fss on fsts.student_session_id = fss.id
    left join feedback_session fs on fs.id = fss.session_id
    left join feedback_course_faculty_template fcft on fcft.id = fsts.feedback_course_faculty_template_id
    left join class cl on cl.id = fcft.class_id
    left join faculty_profile fp on fp.ukid = fcft.faculty_ukid
    left join feedback_template ft on ft.id = fcft.template_id
    left join course c on c.course_id = cl.course_id
    left join programme pr on pr.programme_id = cl.programme_id
    left join term on term.id = cl.term_id
    left join user_attributes ua on ua.ukid = fp.ukid
    where fsr.option_id is null
    and fss.session_id = %d and fsts.submitted = 1
    order by template_name, batch, faculty_name, question_id;
    ''' %x)

    options = fetch_data('''
    Select feedback_question_id, option_text, option_score
    from feedback_template_question_option ftqo
    ''')

    # template definition (sections / questions / options) used to build the
    # objective report directly, instead of loading a .docx template per template.
    template_structure = fetch_data('''
    select t4.id as template_id, t4.template_name, t3.id as section_id, t3.section_name,
           t2.id as question_id, t2.question_text, t2.question_type,
           t1.option_text, t1.option_score
    from feedback_template t4
    left join feedback_template_section t3 on t3.feedback_template_id = t4.id
    left join feedback_template_section_question t2 on t2.template_section_id = t3.id
    left join feedback_template_question_option t1 on t1.feedback_question_id = t2.id
    order by t4.id, t3.id, t2.id, t1.option_score
    ''')

    classStrength = fetch_data('''
    select t2.term_id, t1.faculty_id faculty_ukid, t2.course_id,t2.id as class_id,t2.batch as class_name, count(distinct t4.ukid) studentPresent from class_faculty t1
        left join class t2 on t1.class_id = t2.id
        left join course t3 on t2.course_id = t3.course_id
        left join class_student t4 on t1.class_id = t4.class_id
        group by t1.faculty_id, t2.course_id, t2.term_id,t2.id,t2.batch
    ''')

    # 🧹 Clean up strings to prevent mismatched joins
    for col in ['faculty_name', 'course_code', 'course_name', 'template_name', 'term','class_id','class_name']:
        response[col] = response[col].astype(str).str.strip()
        response_subjective[col] = response_subjective[col].astype(str).str.strip()



    filter1 = response[['session_name', 'template_name', 'term', 'faculty_ukid', 'faculty_name', 'course_id', 'course_code', 'course_name', 'term_id','class_id','class_name','course_type']].drop_duplicates(subset=['faculty_name', 'course_code', 'course_name','template_name','class_id','class_name'])
    templates = filter1['template_name'].drop_duplicates().to_list()

    for template in templates:
        print("Template: ", template)

        # resolve the template id, then its DB definition (sections/questions/options)
        tid_series = response.loc[response['template_name'] == template, 'template_id']
        tid = tid_series.iloc[0] if len(tid_series) else None
        sdf = template_structure[template_structure['template_id'] == tid] if tid is not None else pd.DataFrame()

        if tid is not None and not sdf.empty:

            tstruct = parse_template_structure(sdf)
            obj_question_ids = tstruct['obj_question_ids']

            location = None
            newColsOption = []
            newColsWeight = []
            lst = []
            qweight = []
            questionLen = 0
            totalResponse = []
            sectionLen = 0
            sectionName = ''

            itr = filter1[filter1['template_name'] == template].reset_index(drop=True)
            unique = itr.drop_duplicates(subset=['faculty_name', 'template_name', 'course_code', 'course_name', 'term','class_id','class_name'])

            # loop for generation of excel document to loop through each faculty and course
            # this will act as raw data for the document generation later
            lst = []
            qweight = []
            newColsOption = []
            newColsWeight = []

            for i, r in unique.iterrows():

                d1 = response[(response['template_name'] == r['template_name']) &
                            (response['faculty_name'] == r['faculty_name']) &
                            (response['course_code'] == r['course_code']) &
                            (response['class_id'] == r['class_id']) &
                            (response['term'] == r['term'])].copy()

                d1['option_text'] = d1['option_text'].astype(str).str.strip()

                totalResponse.append(d1['student_ukid'].nunique())

                # drive question order from the DB template structure so that
                # itemRating{n} / Q{n}_* columns line up with the generated table.
                question = obj_question_ids
                questionLen = len(question)

                row_counts = []
                row_weights = []

                for index, q in enumerate(question):

                    d2 = d1[d1['question_id'] == q]

                    opt_df = options[options['feedback_question_id'] == q][['option_text','option_score']].copy()
                    opt_df['option_text'] = opt_df['option_text'].astype(str).str.strip()

                    # aggregate: option -> number of students who selected it
                    agg = (
                        d2.groupby('option_text')['student_ukid']
                        .nunique()
                    )

                    total_selections = int(agg.sum())   # 🔑 NEW

                    if i == 0:
                        for opt in opt_df['option_text']:
                            newColsOption.append(f"Q{index+1}_{opt}")
                        newColsWeight.append(f"W{index+1}")
                        newColsWeight.append(f"C{index+1}")   # 🔑 NEW (count column)

                    qscore = 0
                    for _, opt in opt_df.iterrows():
                        text = opt['option_text']
                        score = opt['option_score']

                        count = int(agg.get(text, 0))
                        row_counts.append(count)
                        qscore += count * score

                    row_weights.append(qscore)
                    row_weights.append(total_selections)   # 🔑 NEW


                lst.append(row_counts)
                qweight.append(row_weights)


            
            optionDF = pd.DataFrame(lst, columns=newColsOption)
            weightDF = pd.DataFrame(qweight, columns=newColsWeight)


            finalTemp = pd.merge(unique,classStrength, left_on=['term_id', 'faculty_ukid', 'course_id','class_name'], right_on=['term_id', 'faculty_ukid', 'course_id','class_name'], how='left')

            finalTemp.drop(columns=['class_id_x'], inplace=True)
            finalTemp.rename(columns={'class_id_y': 'class_id'}, inplace=True)
            finalDataset = pd.concat([finalTemp, pd.DataFrame({'totalResponse': totalResponse}), optionDF, weightDF], axis=1)
            
            # check if output folder exists, if not create it
            location = folder_check(r['term'] + '-' + template)

            for eachQuestion in range(questionLen):
                weight_col = f"W{eachQuestion+1}"
                count_col  = f"C{eachQuestion+1}"

                # guard against questions with no responses (count == 0)
                denom = finalDataset[count_col].replace(0, np.nan)
                finalDataset[f'itemRating{eachQuestion+1}'] = (finalDataset[weight_col] / denom).round(2)

            # ---- section averages derived from the DB template structure ----
            sec_cols = []
            for k, sec in enumerate(tstruct['sections'], start=1):
                item_cols = [f"itemRating{q['position']}" for q in sec['questions']]
                if not item_cols:
                    continue
                finalDataset[f'sec{k}'] = finalDataset[item_cols].mean(axis=1).round(2)
                sec_cols.append(f'sec{k}')

            item_cols_all = [f'itemRating{i+1}' for i in range(questionLen)]
            finalDataset['final'] = finalDataset[item_cols_all].mean(axis=1).round(2)
            if len(sec_cols) > 1:
                finalDataset['section_final'] = finalDataset[sec_cols].mean(axis=1).round(2)


            finalDataset.to_excel(str(run_root() / (r['term'] + '-' + template + ".xlsx")), index=False)
            # generation of subjective and objective documents
            # create subjective document
            for i1, r1 in unique.iterrows():
                d1 = response[(response['template_name'] == r1['template_name']) &
                            (response['faculty_name'] == r1['faculty_name']) &
                            (response['course_code'] == r1['course_code']) &
                            (response['class_id'] == r1['class_id']) &
                            (response['term'] == r1['term'])]
                
                subDocument = createSubjectiveDocument(dataSub=response_subjective[(response_subjective['template_name'] == r1['template_name']) &
                        (response_subjective['faculty_name'] == r1['faculty_name']) &
                        (response_subjective['course_code'] == r1['course_code']) &
                        (response_subjective['class_id'] == r1['class_id']) &
                        (response_subjective['class_name'] == r1['class_name']) &
                        (response_subjective['term'] == r1['term'])], 
                        faculty=r1['faculty_name'], 
                        course=r1['course_name'], 
                        term=r1['term'], 
                        course_code=r1['course_code'], 
                        # class_id=r1['class_id'],
                        class_name=r1['class_name'],
                        saveLoc=location)

            # create objective document directly from the DB template structure
            for index, data in finalDataset.iterrows():
                build_objective_document(tstruct, data, template, logo_path, location)


            # loop to merge the objective and subjective documents
            for i2, r2 in unique.iterrows():
                try:
                    # Sanitize course_code for safe filenames
                    safe_code = safe_filename(r2['course_code'][0:13])
                    safe_class   = safe_filename(r2['class_name'])
                    fname = build_report_name(
                        "",
                        r2['faculty_name'],
                        r2['course_code'],
                        r2['class_name']
                    ).lstrip("-")


                    pdf_writer = PdfWriter()

                    # List of PDFs to merge
                    pdf_files = [
                        location / f"obj-{fname}.pdf",
                        location / f"sub-{fname}.pdf"
                    ]

                    # Only keep existing, non-empty PDFs
                    valid_files = [f for f in pdf_files if f.exists() and f.stat().st_size > 0]

                    if not valid_files:
                        print(f"⚠️ Skipping merge for {fname}, no valid PDFs found")
                        continue

                    for pdf_file in valid_files:
                        pdf_reader = PdfReader(pdf_file)
                        for page in range(len(pdf_reader.pages)):
                            pdf_writer.add_page(pdf_reader.pages[page])

                    # Write merged file
                    out_file = location / f"{fname}.pdf"
                    with open(out_file, "wb") as output_pdf:
                        pdf_writer.write(output_pdf)

                    # Clean up temporary files
                    for f in valid_files:
                        try:
                            os.remove(f)
                        except Exception as e:
                            print(f"⚠️ Could not delete {f}: {e}")

                    print(f"✅ Merged report created: {out_file}")

                except Exception as e:
                    print(f"❌ Error merging for {r2['faculty_name']} / {r2['course_code']}: {e}")
                    continue
        
        else:
            print("No DB template definition found for: ", template)
            continue
        # break

